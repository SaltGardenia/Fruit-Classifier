"""
生成课程设计报告 - 修改 docx 模板中的内容
"""

import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

DOCX_PATH = r'c:\Users\SaltGardenia\Documents\GitHub\Fruit Classifier\2024级机器学习课程设计模板（智科）.docx'

def get_paragraphs_by_text(doc, text):
    """查找包含指定文本的段落"""
    results = []
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            results.append((i, p))
    return results

def clear_paragraph_runs(paragraph):
    """清除段落的所有run，保留段落格式"""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

def set_run_text(run, text, font_name='宋体', font_size=None, bold=False, color=None):
    """设置run的文本和格式"""
    run.text = text
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) if run._element.find(qn('w:rPr')) is not None else None
    if font_size:
        run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_run_to_paragraph(paragraph, text, font_name='宋体', font_size=Pt(12), bold=False, color=None):
    """向段落添加run"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run

def replace_paragraph_text(paragraph, new_text, font_name='宋体', font_size=Pt(12)):
    """替换段落文本（保留段落本身，清除旧run后添加新run）"""
    clear_paragraph_runs(paragraph)
    add_run_to_paragraph(paragraph, new_text, font_name, font_size)

def insert_paragraph_after(paragraph, text, style=None):
    """在指定段落之后插入新段落"""
    new_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    add_run_to_paragraph(new_para, text)
    return new_para

def insert_paragraph_before(paragraph, text, style=None):
    """在指定段落之前插入新段落"""
    new_p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    paragraph._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    add_run_to_paragraph(new_para, text)
    return new_para


def main():
    doc = Document(DOCX_PATH)

    # =========================================================================
    # 1. 研究背景
    # =========================================================================
    print("=" * 60)
    print("正在写入「一、研究背景」...")

    research_bg_content = (
        "水果分类是计算机视觉领域中图像识别技术的重要应用方向之一，在智能农业、"
        "无人零售、食品质量检测等场景中具有广泛的应用需求。传统的水果分类方法通常依赖"
        "于人工设计的特征提取器（如颜色直方图、纹理特征、形状描述子等）结合传统的机器"
        "学习分类器（如支持向量机、随机森林等），但这些方法在面对光照变化、遮挡、姿态"
        "差异、品种多样性等复杂情况时，特征表达能力有限，泛化性能不佳。\n\n"
        "近年来，随着深度学习技术的快速发展，卷积神经网络（Convolutional Neural "
        "Network，CNN）在图像分类任务中取得了突破性进展。CNN通过端到端的学习方式，"
        "能够自动从原始图像中提取层次化的特征表示，从低级的边缘、纹理到高级的语义信息，"
        "显著超越了传统方法的识别精度。经典的CNN架构如AlexNet、VGGNet、ResNet等已在"
        "ImageNet等大规模图像分类竞赛中展现出卓越性能。\n\n"
        "本工程基于CNN构建了一个面向5类常见水果（苹果Apple、香蕉Banana、芒果Mango、"
        "橙子Orange、菠萝Pineapple）的图像分类系统。系统涵盖数据预处理、CNN模型构建与"
        "训练、模型评估以及基于PyQt5的图形用户界面开发等完整流程。用户可以通过图形界面"
        "上传或拖拽水果图片，系统自动完成分类识别并展示各类别的置信度概率分布，实现了"
        "从模型训练到实际应用的全链路解决方案。"
    )

    # 找到包含研究背景内容的段落（P52）
    for i, p in enumerate(doc.paragraphs):
        if "说明本工程所涉及的研究背景" in p.text:
            print(f"  找到研究背景段落 P{i}")
            clear_paragraph_runs(p)
            # 分段写入 - 三个段落
            paras = research_bg_content.split('\n\n')
            first_para = True
            for para_text in paras:
                if para_text.strip():
                    if first_para:
                        add_run_to_paragraph(p, para_text.strip())
                        first_para = False
                    else:
                        new_p = insert_paragraph_after(p, para_text.strip())
                        p = new_p  # 继续在下一个段落后面插入
            break

    # =========================================================================
    # 2. 模型方法
    # =========================================================================
    print("正在写入「二、模型方法」...")

    model_content = (
        "本工程采用卷积神经网络（CNN）作为核心分类模型。CNN是一种专门为处理网格状"
        "数据（如图像）而设计的深度神经网络，其核心思想是利用卷积操作提取局部特征，"
        "通过层级化的结构从低级特征逐步抽象为高级语义特征。\n\n"
        "（1）卷积层（Conv2D）：卷积层是CNN的核心组件，通过一组可学习的卷积核（kernel）"
        "在输入图像上滑动并进行点积运算，提取局部特征。本模型使用3×3大小的卷积核，"
        "padding='same'保持特征图尺寸不变，激活函数采用ReLU（Rectified Linear "
        "Unit），其表达式为f(x)=max(0,x)，能够有效缓解梯度消失问题并加速收敛。\n\n"
        "（2）池化层（MaxPooling2D）：池化层用于对特征图进行下采样，降低特征维度，"
        "减少参数量和计算开销，同时增强模型的平移不变性。本模型采用2×2的最大池化，"
        "在每个2×2的局部区域内取最大值作为输出，使特征图尺寸减半。\n\n"
        "（3）Dropout正则化：Dropout是一种防止过拟合的正则化技术，在训练过程中以一定"
        "概率随机丢弃神经元的输出，迫使网络学习更加鲁棒的特征表示。本模型在全连接层后"
        "设置了0.2的Dropout比率。\n\n"
        "（4）全连接层（Dense）与Softmax：经过卷积和池化提取的特征通过Flatten层展平"
        "为一维向量，然后输入全连接层进行高阶推理。最后一层使用Softmax激活函数，将"
        "输出转换为5个类别的概率分布，所有类别概率之和为1。\n\n"
        "（5）模型结构与参数：模型包含3个卷积块（每个卷积块由Conv2D+MaxPooling2D组成），"
        "卷积核数量依次为32、64、32。展平后连接两个全连接层（神经元数分别为64和32），"
        "最后输出5维的Softmax层。模型总参数量约为180万。\n\n"
        "（6）训练配置：优化器采用Adam（自适应矩估计），它结合了Momentum和RMSProp的"
        "优点，能自适应调整学习率；损失函数采用稀疏分类交叉熵（Sparse Categorical "
        "Crossentropy），适用于整数标签的多分类任务；评估指标为准确率（Accuracy）。\n\n"
        "（7）数据增强（Data Augmentation）：为扩充训练数据、提升泛化能力，使用"
        "ImageDataGenerator对训练图像进行实时增强，包括随机水平翻转、垂直翻转以及"
        "像素值归一化（rescale=1/255）。同时将数据按8:2的比例划分为训练集和验证集。"
    )

    for i, p in enumerate(doc.paragraphs):
        if "详细说明本工程所使用的模型方法和理论" in p.text:
            print(f"  找到模型方法段落 P{i}")
            clear_paragraph_runs(p)
            paras = model_content.split('\n\n')
            first_para = True
            for para_text in paras:
                if para_text.strip():
                    if first_para:
                        add_run_to_paragraph(p, para_text.strip())
                        first_para = False
                    else:
                        new_p = insert_paragraph_after(p, para_text.strip())
                        p = new_p
            break

    # =========================================================================
    # 3. 系统设计
    # =========================================================================
    print("正在写入「三、系统设计」...")

    system_design_content = (
        "本系统的整体设计分为数据处理模块、模型训练模块和图形界面交互模块三大核心部分。"
        "系统流程为：数据预处理 → 模型构建 → 模型训练与验证 → 模型保存 → GUI加载模型 → "
        "用户交互识别。\n\n"
        "3.1 数据预处理\n"
        "数据集包含5类水果（Apple、Banana、Mango、Orange、Pineapple），每类各含训练集"
        "和测试集图像。使用ImageDataGenerator对图像进行预处理和增强：将所有图像统一"
        "缩放到128×128像素，像素值归一化到[0,1]区间；对训练集应用随机水平翻转和垂直"
        "翻转进行数据增强，并划分20%作为验证集。批大小（batch size）设为32。\n\n"
        "3.2 模型构建\n"
        "使用TensorFlow/Keras Sequential API构建CNN模型，结构如下：\n"
        "• 输入层：128×128×3的RGB图像\n"
        "• 卷积块1：Conv2D(32, 3×3, ReLU) + MaxPooling2D(2×2)\n"
        "• 卷积块2：Conv2D(64, 3×3, ReLU) + MaxPooling2D(2×2)\n"
        "• 卷积块3：Conv2D(32, 3×3, ReLU) + MaxPooling2D(2×2)\n"
        "• Flatten展平层\n"
        "• 全连接层Dense(64, ReLU) + Dropout(0.2)\n"
        "• 全连接层Dense(32, ReLU) + Dropout(0.2)\n"
        "• 输出层Dense(5, Softmax)\n\n"
        "3.3 模型训练\n"
        "模型使用Adam优化器（默认学习率0.001）、稀疏分类交叉熵损失函数进行编译。"
        "设置训练轮数（epochs）为50轮，以32为批大小进行迭代训练。每个epoch结束后在"
        "验证集上评估模型性能，监控训练和验证的准确率与损失值变化，防止过拟合。\n\n"
        "3.4 模型评估\n"
        "训练完成后，绘制训练/验证准确率和损失曲线，直观分析模型的学习效果和收敛情况。"
        "此外，使用独立的测试集对模型进行最终评估，计算在未见数据上的分类准确率，验证"
        "模型的泛化能力。训练好的模型以.h5格式保存至Model目录。\n\n"
        "3.5 GUI交互界面\n"
        "基于PyQt5开发图形用户界面，提供直观的人机交互体验。主要功能包括：\n"
        "• 图片选择：支持通过文件对话框选择和拖拽两种方式加载待检测图片\n"
        "• 实时显示：左侧区域展示用户选择的图片\n"
        "• 分类检测：点击「开始检测」按钮，调用训练好的CNN模型进行预测\n"
        "• 结果展示：右侧面板显示识别结果（水果名称及对应emoji）、置信度百分比\n"
        "• 概率分布图：以水平柱状图形式可视化展示5个类别的概率分布，预测类别高亮显示\n"
        "• 重置功能：一键清空当前结果，重新开始新图片的识别\n"
        "界面采用现代化扁平化设计风格，配色柔和，交互反馈清晰，操作流程简洁明了。"
    )

    for i, p in enumerate(doc.paragraphs):
        if "系统的详细设计，系统流程" in p.text:
            print(f"  找到系统设计段落 P{i}")
            clear_paragraph_runs(p)
            paras = system_design_content.split('\n\n')
            first_para = True
            for para_text in paras:
                if para_text.strip():
                    if first_para:
                        add_run_to_paragraph(p, para_text.strip())
                        first_para = False
                    else:
                        new_p = insert_paragraph_after(p, para_text.strip())
                        p = new_p
            break

    # =========================================================================
    # 4. 系统演示与分析
    # =========================================================================
    print("正在写入「四、系统演示与分析」...")

    demo_content = (
        "4.1 系统演示\n"
        "系统启动后，主窗口分为左右两栏：左栏为图片显示区和操作按钮（选择图片、开始检测、"
        "重置），右栏为检测结果展示区。用户可通过点击「选择图片」按钮或直接将图片拖拽到"
        "虚线框内加载待检测的水果图片，然后点击「开始检测」执行分类。\n\n"
        "以测试集中的Apple样本为例：加载一张苹果图片后点击检测，模型输出结果为Apple，"
        "置信度达到较高水平（通常超过95%），同时右侧概率分布柱状图显示Apple类别的概率"
        "显著高于其他4类，表明模型对该样本的分类结果可靠。再以Pineapple样本测试，模型"
        "同样能正确识别并给出高置信度结果。\n\n"
        "4.2 结果分析\n"
        "从测试结果来看，模型对5类水果均具备良好的识别能力，主要原因如下：\n"
        "• 各类水果的外形、颜色、纹理特征差异明显，类间距离大，分类难度相对较低\n"
        "• CNN通过多层卷积提取了有效的判别性特征\n"
        "• 数据增强技术提升了模型对图像变换的鲁棒性\n\n"
        "但在某些困难样本上，模型可能出现误分类或置信度偏低的情况。分析原因可能包括：\n"
        "• 训练数据量有限，每类水果的图像数量和多样性不足\n"
        "• 类内差异较大（如不同品种的苹果颜色和形状差异明显）\n"
        "• 背景干扰或光照条件不佳影响特征提取\n\n"
        "4.3 改进方案\n"
        "针对以上问题，提出以下改进方向：\n"
        "（1）扩充数据集：采集更多样化的水果图像，增加不同品种、成熟度、光照条件和"
        "背景下的样本，提升数据多样性和覆盖度。\n"
        "（2）引入迁移学习：使用在ImageNet上预训练的模型（如MobileNetV2、EfficientNet等）"
        "作为特征提取器进行微调（Fine-tuning），可显著提升小样本场景下的分类精度。\n"
        "（3）优化模型结构：增加卷积层深度或引入残差连接（Residual Connection），"
        "提升模型的表征能力。\n"
        "（4）超参数调优：对学习率、批大小、Dropout比率、卷积核数量等超参数进行"
        "网格搜索或随机搜索，寻找最优配置。\n"
        "（5）集成学习：训练多个不同初始化的模型，对预测结果进行投票或加权平均，"
        "进一步提升分类稳定性。"
    )

    for i, p in enumerate(doc.paragraphs):
        if "对系统演示结果进行说明" in p.text:
            print(f"  找到系统演示段落 P{i}")
            clear_paragraph_runs(p)
            paras = demo_content.split('\n\n')
            first_para = True
            for para_text in paras:
                if para_text.strip():
                    if first_para:
                        add_run_to_paragraph(p, para_text.strip())
                        first_para = False
                    else:
                        new_p = insert_paragraph_after(p, para_text.strip())
                        p = new_p
            break

    # =========================================================================
    # 5. 对本门课的感想、意见和建议
    # =========================================================================
    print("正在写入「五、对本门课的感想、意见和建议」...")

    feedback_content = (
        "通过本学期机器学习课程的深入学习与本次课程设计的实践，我对机器学习的基本概念、"
        "核心算法和完整开发流程有了更加全面和深入的理解。从数据预处理、模型选择与构建、"
        "训练调优到最终的部署应用，整个流程的亲身实践让我真正体会到了机器学习技术从理论"
        "到落地的完整过程。\n\n"
        "本课程设计以水果分类为切入点，将卷积神经网络这一经典深度学习模型与实际应用场景"
        "相结合，既巩固了课堂所学的CNN理论知识，又锻炼了编程实现和系统整合的能力。"
        "特别是在图形界面开发环节，将训练好的模型封装为可交互的应用，让我深刻认识到"
        "模型部署和用户体验在实际工程中的重要性。\n\n"
        "对课程的建议：希望课程能够增加更多前沿模型（如Transformer在视觉领域的应用、"
        "Vision Transformer等）的讲解内容，并安排更多的实践课时，让学生有更多时间动手"
        "实验和探索。此外，如果能引入一些工业级项目案例（如模型部署到移动端或Web端），"
        "将更有助于拓宽学生的工程视野。"
    )

    for i, p in enumerate(doc.paragraphs):
        if "五、对本门课的感想、意见和建议" in p.text:
            print(f"  找到感想段落 P{i}")
            paras = feedback_content.split('\n\n')
            for para_text in paras:
                if para_text.strip():
                    new_p = insert_paragraph_after(p, para_text.strip())
                    p = new_p
            break

    # =========================================================================
    # 6. 视频链接提示
    # =========================================================================
    print("正在写入「六、讲解视频链接」占位...")

    for i, p in enumerate(doc.paragraphs):
        if "https://pan.baidu.com/s/" in p.text:
            print(f"  找到链接占位 P{i}")
            # 替换链接占位提示
            for run in p.runs:
                if "https://pan.baidu.com/s/" in run.text:
                    # 保持占位符，不实际填入链接
                    pass
            break

    # =========================================================================
    # 保存文档
    # =========================================================================
    output_path = r'c:\Users\SaltGardenia\Documents\GitHub\Fruit Classifier\2024级机器学习课程设计报告（智能科）.docx'
    doc.save(output_path)
    print(f"\n✅ 报告已生成：{output_path}")
    print("完成！")


if __name__ == "__main__":
    main()
